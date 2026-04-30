# -*- coding: utf-8 -*-
"""
Build letterhead-thiqqah.docx — 3-column header (EN | AR | logo), footer QR line.
Creates letterhead-watermark-print.png with Pillow. On Windows + Word + pywin32, embeds
that image as a centered behind-text watermark after save.

مصدر البيانات الوحيد: thiqqah-site/thiqqah-brand.json
يولّد أيضاً: thiqqah-brand.embed.js + assets/letterhead-qr.png (للصفحة المطبوعة).
"""
from __future__ import annotations

import io
import json
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[2]
BRAND_JSON = REPO / "thiqqah-site" / "thiqqah-brand.json"
EMBED_JS_OUT = REPO / "thiqqah-site" / "thiqqah-brand.embed.js"
OUT_PATH = REPO / "thiqqah-site" / "letterhead-thiqqah.docx"

CLR_INK = RGBColor(0x12, 0x24, 0x3A)
CLR_GREEN = RGBColor(0x07, 0x3F, 0x3A)
CLR_GREEN2 = RGBColor(0x0A, 0x66, 0x5D)
CLR_MUTED = RGBColor(0x5F, 0x6D, 0x7F)


def load_brand() -> dict:
    if not BRAND_JSON.is_file():
        raise FileNotFoundError(f"ضع بيانات الهوية في: {BRAND_JSON}")
    data = json.loads(BRAND_JSON.read_text(encoding="utf-8"))
    for key in (
        "siteUrl",
        "displayUrl",
        "email",
        "phoneE164",
        "phoneDisplay",
        "commercialRegistration",
        "cityLineAr",
        "nameAr",
        "nameEn",
        "taglineEnTemplate",
        "logoRelativePath",
        "watermarkRelativePath",
        "qrImageRelativePath",
    ):
        if key not in data:
            raise KeyError(f"thiqqah-brand.json ناقص المفتاح: {key}")
    ne = data["nameEn"]
    for nk in ("line1", "line2", "line3"):
        if nk not in ne:
            raise KeyError(f"thiqqah-brand.json nameEn ناقص: {nk}")
    return data


def write_embed_js(brand: dict) -> None:
    lines = [
        "// تُولَّد تلقائياً — لا تعدّل يدوياً.",
        "// المصدر الوحيد: thiqqah-brand.json ← ثم شغّل RUN-LETTERHEAD.cmd",
        "window.__THIQQAH_BRAND = ",
        json.dumps(brand, ensure_ascii=False, indent=2),
        ";\n",
    ]
    EMBED_JS_OUT.write_text("\n".join(lines), encoding="utf-8")
    print("Wrote:", EMBED_JS_OUT)


def write_qr_png(url: str, dest: Path) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(qr_png_bytes(url))
    print("Wrote:", dest)


def en_block_from_brand(brand: dict) -> str:
    cr = brand["commercialRegistration"]
    tag = brand["taglineEnTemplate"].replace("{cr}", cr)
    ne = brand["nameEn"]
    return f"{ne['line1']}\n{ne['line2']}\n{ne['line3']}\n{tag}"


def set_paragraph_rtl(paragraph) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def format_ar_paragraph(paragraph, *, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.alignment = align
    if rtl:
        set_paragraph_rtl(paragraph)


def run_no_spellcheck(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    r_pr.append(OxmlElement("w:noProof"))


def set_cell_margins_twips(cell, *, top=0, left=0, bottom=0, right=0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, val in (("w:top", top), ("w:left", left), ("w:bottom", bottom), ("w:right", right)):
        el = OxmlElement(tag)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def add_page_border(section) -> None:
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "double")
        el.set(qn("w:sz"), "22")
        el.set(qn("w:space"), "8")
        el.set(qn("w:color"), "073F3A")
        pg_borders.append(el)
    sect_pr.append(pg_borders)


def qr_png_bytes(url: str) -> bytes:
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=5,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#073f3a", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def write_soft_watermark_png(src: Path, dst: Path) -> bool:
    try:
        from PIL import Image
    except ImportError:
        return False
    im = Image.open(src).convert("RGBA")
    bands = im.split()
    if len(bands) < 4:
        im = im.convert("RGBA")
        bands = im.split()
    alpha = bands[3]
    im.putalpha(alpha.point(lambda p: int(p * 0.11)))
    dst.parent.mkdir(parents=True, exist_ok=True)
    im.save(dst, format="PNG")
    return True


def embed_watermark_word_com(docx_path: Path, png_path: Path) -> bool:
    """Insert faint PNG as page-centered behind-text shape in primary header (Word on Windows)."""
    docx_path = docx_path.resolve()
    png_path = png_path.resolve()
    if not docx_path.is_file() or not png_path.is_file():
        return False
    try:
        import win32com.client as win32
    except ImportError:
        print("Optional: pip install pywin32  (auto-embed watermark in Word)")
        return False

    wd_header_primary = 1
    wd_wrap_behind = 7
    wd_rel_page = 1

    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    ok = False
    try:
        doc = word.Documents.Open(
            str(docx_path), ReadOnly=False, AddToRecentFiles=False
        )
        sec = doc.Sections(1)
        hdr = sec.Headers(wd_header_primary)
        for idx in range(int(hdr.Shapes.Count), 0, -1):
            try:
                if hdr.Shapes(idx).Name == "ThiqqahWaterMark":
                    hdr.Shapes(idx).Delete()
            except Exception:
                pass

        w_pt = 240.0
        shape = hdr.Shapes.AddPicture(
            str(png_path), False, True, 0, 0, w_pt, w_pt
        )
        shape.Name = "ThiqqahWaterMark"
        shape.WrapFormat.Type = wd_wrap_behind
        shape.RelativeHorizontalPosition = wd_rel_page
        shape.RelativeVerticalPosition = wd_rel_page
        try:
            shape.LockAnchor = True
        except Exception:
            pass
        ph = float(sec.PageSetup.PageHeight)
        pw = float(sec.PageSetup.PageWidth)
        shape.Left = (pw - float(shape.Width)) / 2.0
        shape.Top = (ph - float(shape.Height)) / 2.0
        try:
            shape.PictureFormat.Brightness = 0.9
            shape.PictureFormat.Contrast = 0.15
        except Exception:
            pass

        doc.Save()
        doc.Close()
        ok = True
        print("Embedded watermark in Word (COM).")
    except Exception as e:
        print("Word COM:", type(e).__name__, "-", str(e)[:120])
    finally:
        try:
            word.Quit()
        except Exception:
            pass
    return ok


def main() -> None:
    brand = load_brand()
    write_embed_js(brand)
    site_root = REPO / "thiqqah-site"
    qr_rel = Path(brand["qrImageRelativePath"])
    write_qr_png(brand["siteUrl"], site_root / qr_rel)

    logo_primary = site_root / Path(brand["logoRelativePath"])
    fallback_logo = site_root / "assets" / "thiqqah-logo.png"
    logo = logo_primary if logo_primary.is_file() else fallback_logo
    if not logo.is_file():
        raise FileNotFoundError(f"الشعار غير موجود: {logo_primary} أو {fallback_logo}")

    wm_src = site_root / Path(brand["watermarkRelativePath"])
    if not wm_src.is_file():
        wm_src = logo_primary if logo_primary.is_file() else logo
    WM_OUT = site_root / "assets" / "letterhead-watermark-print.png"
    has_wm_png = write_soft_watermark_png(wm_src, WM_OUT)
    if has_wm_png:
        print("Watermark PNG:", WM_OUT)

    doc = Document()
    sect = doc.sections[0]
    sect.page_height = Cm(29.7)
    sect.page_width = Cm(21.0)
    sect.top_margin = Cm(2.0)
    sect.bottom_margin = Cm(3.2)
    sect.left_margin = Cm(2.0)
    sect.right_margin = Cm(2.0)
    sect.header_distance = Cm(0.65)
    sect.footer_distance = Cm(0.6)
    add_page_border(sect)

    header = sect.header
    ht = header.add_table(1, 3, Cm(17.0))
    ht.columns[0].width = Cm(5.35)
    ht.columns[1].width = Cm(7.35)
    ht.columns[2].width = Cm(4.3)
    c_en, c_ar, c_logo = ht.rows[0].cells[0], ht.rows[0].cells[1], ht.rows[0].cells[2]
    for c in (c_en, c_ar, c_logo):
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    set_cell_margins_twips(c_logo, top=0, left=40, bottom=0, right=0)
    set_cell_margins_twips(c_en, top=0, left=0, bottom=0, right=80)
    set_cell_margins_twips(c_ar, top=0, left=80, bottom=0, right=80)

    # ——— عمود إنجليزي ———
    p_en = c_en.paragraphs[0]
    p_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_en.paragraph_format.space_after = Pt(0)
    r_en = p_en.add_run(en_block_from_brand(brand))
    r_en.font.size = Pt(8.5)
    r_en.font.name = "Times New Roman"
    r_en.font.color.rgb = CLR_GREEN2
    run_no_spellcheck(r_en)

    # ——— عمود عربي ———
    p1 = c_ar.paragraphs[0]
    format_ar_paragraph(p1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p1.paragraph_format.space_after = Pt(1)
    nar = brand["nameAr"]
    r1 = p1.add_run(f"{nar['beforeGold']}{nar['gold']}\n{nar['line2']}")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.name = "Arial"
    r1.font.color.rgb = CLR_GREEN
    run_no_spellcheck(r1)

    p2 = c_ar.add_paragraph()
    format_ar_paragraph(p2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(f"السجل التجاري: {brand['commercialRegistration']}")
    r2.font.size = Pt(9)
    r2.font.name = "Arial"
    r2.bold = True
    r2.font.color.rgb = CLR_INK
    run_no_spellcheck(r2)

    p3 = c_ar.add_paragraph()
    format_ar_paragraph(p3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r3 = p3.add_run(brand["cityLineAr"])
    r3.font.size = Pt(8.5)
    r3.font.name = "Arial"
    r3.font.color.rgb = CLR_MUTED
    run_no_spellcheck(r3)

    # ——— عمود الشعار ———
    pl = c_logo.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(pl)
    pl.paragraph_format.space_after = Pt(0)
    pl.add_run().add_picture(str(logo), width=Cm(3.35))

    rule = header.add_table(1, 1, Cm(17.0))
    rc = rule.rows[0].cells[0]
    tcp = rc._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "20")
    bot.set(qn("w:color"), "BD9144")
    borders.append(bot)
    tcp.append(borders)
    rc.paragraphs[0].paragraph_format.space_after = Pt(2)

    for _ in range(22):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = Pt(14)

    footer = sect.footer
    sep = footer.add_table(1, 1, Cm(17.0))
    sc = sep.rows[0].cells[0]
    stc = sc._tc.get_or_add_tcPr()
    sb = OxmlElement("w:tcBorders")
    st = OxmlElement("w:top")
    st.set(qn("w:val"), "single")
    st.set(qn("w:sz"), "24")
    st.set(qn("w:color"), "BD9144")
    sb.append(st)
    stc.append(sb)
    sc.paragraphs[0].paragraph_format.space_after = Pt(1)

    row = footer.add_table(1, 2, Cm(17.0))
    row.columns[0].width = Cm(13.9)
    row.columns[1].width = Cm(3.1)
    c_line, c_qr = row.rows[0].cells[0], row.rows[0].cells[1]
    c_line.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    c_qr.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    fp = c_line.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fp.paragraph_format.space_after = Pt(0)
    for txt, bold in [
        (brand["displayUrl"], True),
        ("   ·   ", False),
        (brand["email"], True),
        ("   ·   ", False),
        (brand["phoneDisplay"], True),
    ]:
        rr = fp.add_run(txt)
        rr.font.size = Pt(8.5)
        rr.font.name = "Arial"
        rr.font.color.rgb = CLR_GREEN2
        rr.bold = bold

    pq = c_qr.paragraphs[0]
    pq.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pq.paragraph_format.space_after = Pt(0)
    pq.add_run().add_picture(io.BytesIO(qr_png_bytes(brand["siteUrl"])), width=Cm(1.45))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    saved: Path = OUT_PATH
    try:
        doc.save(str(OUT_PATH))
        print("Saved:", OUT_PATH)
    except PermissionError:
        saved = OUT_PATH.with_name("letterhead-thiqqah-new.docx")
        doc.save(str(saved))
        print("Permission: saved alternate:", saved)

    if has_wm_png and WM_OUT.is_file():
        embed_watermark_word_com(saved, WM_OUT)


if __name__ == "__main__":
    main()
